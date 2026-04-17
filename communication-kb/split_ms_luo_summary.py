"""
Split 摩根士丹利罗总2025-2026年交流记录汇总.docx into one .docx per session.

Sessions are delimited by a standalone paragraph matching YYYY/M/D (or MM/DD with year).
Output files: 摩根士丹利罗总线下沟通YYYYMMDD.docx
- 2025-dated sessions -> Desktop/2025交流记录
- 2026-dated sessions -> Desktop/2026交流记录

The original file is moved to Desktop/2026交流记录/_archive_已拆分/ (not scanned by ingest).

Run: py -3 communication-kb/split_ms_luo_summary.py
"""
from __future__ import annotations

import re
import shutil
from pathlib import Path

from docx import Document

DATE_LINE = re.compile(r"^(\d{4})/(\d{1,2})/(\d{1,2})$")
SOURCE_NAME = "摩根士丹利罗总2025-2026年交流记录汇总.docx"


def ymd_from_match(m: re.Match[str]) -> str:
    y, mo, d = m.group(1), int(m.group(2)), int(m.group(3))
    return f"{y}{mo:02d}{d:02d}"


def desk() -> Path:
    return Path.home() / "Desktop"


def archive_dir() -> Path:
    p = desk() / "2026交流记录" / "_archive_已拆分"
    p.mkdir(parents=True, exist_ok=True)
    return p


def dest_root_for_ymd(ymd: str) -> Path:
    y = ymd[:4]
    folder = desk() / f"{y}交流记录"
    folder.mkdir(parents=True, exist_ok=True)
    return folder


def split_paragraphs(lines: list[str]) -> tuple[str, list[tuple[str, str, list[str]]]]:
    """Return (title, [(ymd, date_line, body_lines), ...])."""
    if not lines:
        return ("", [])
    i = 0
    title = "摩根士丹利罗总交流纪要"
    if not DATE_LINE.match(lines[0]):
        title = lines[0]
        i = 1
    chunks: list[tuple[str, str, list[str]]] = []
    while i < len(lines):
        if not DATE_LINE.match(lines[i]):
            i += 1
            continue
        date_line = lines[i]
        m = DATE_LINE.match(date_line)
        assert m
        ymd = ymd_from_match(m)
        i += 1
        body: list[str] = []
        while i < len(lines) and not DATE_LINE.match(lines[i]):
            body.append(lines[i])
            i += 1
        chunks.append((ymd, date_line, body))
    return title, chunks


def write_session_doc(
    out_path: Path,
    *,
    ymd: str,
    date_line: str,
    body: list[str],
) -> None:
    doc = Document()
    doc.add_heading(f"摩根士丹利罗总线下沟通（{ymd}）", level=1)
    doc.add_paragraph(date_line)
    for para in body:
        doc.add_paragraph(para)
    doc.save(str(out_path))


def main() -> None:
    src = desk() / "2026交流记录" / SOURCE_NAME
    if not src.is_file():
        raise SystemExit(f"未找到源文件: {src}")

    doc = Document(str(src))
    lines: list[str] = []
    for p in doc.paragraphs:
        t = (p.text or "").strip()
        if t:
            lines.append(t)

    _title, chunks = split_paragraphs(lines)
    if not chunks:
        raise SystemExit("未解析到任何「YYYY/M/D」日期分段，放弃拆分。")

    written: list[Path] = []
    for ymd, date_line, body in chunks:
        root = dest_root_for_ymd(ymd)
        out_name = f"摩根士丹利罗总线下沟通{ymd}.docx"
        out_path = root / out_name
        if out_path.is_file():
            raise SystemExit(f"目标已存在，停止以免覆盖: {out_path}")
        write_session_doc(out_path, ymd=ymd, date_line=date_line, body=body)
        written.append(out_path)

    arch = archive_dir() / SOURCE_NAME
    if arch.is_file():
        arch.unlink()
    shutil.move(str(src), str(arch))
    print(f"已拆分 {len(chunks)} 场沟通，写入:")
    for p in written:
        print(" ", p)
    print("原文已移至:", arch)


if __name__ == "__main__":
    main()
