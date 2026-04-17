"""
Ingest 交流纪要（默认合并桌面「2026交流记录」「2025交流记录」等）到 kb-data.json + panel.html。
Run from repo root: py -3 communication-kb/ingest.py
"""
from __future__ import annotations

import json
import os
import re
from dataclasses import dataclass, asdict
from datetime import datetime
from pathlib import Path
from typing import Any

from quote_name_resolver import FUND_INDEX_KEY, is_cn_ashare_six, normalize_stock_key

# 非股票六位数字（误识别等），不参与标的解析
STOCK_NUMERIC_BLOCKLIST: frozenset[str] = frozenset({"850000"})


def _default_comm_input_dirs() -> list[Path]:
    """桌面纪要目录：依次尝试 2026、2025、旧名「交流记录」，存在的全部纳入。"""
    desk = Path(os.environ.get("USERPROFILE", str(Path.home()))) / "Desktop"
    candidates = [
        desk / "2026交流记录",
        desk / "2025交流记录",
        desk / "交流记录",
    ]
    return [p.resolve() for p in candidates if p.is_dir()]


def resolve_input_dirs(cli: Path | None) -> list[Path]:
    """--input 单目录优先；否则 COMM_KB_INPUTS（分号分隔多目录）；再否则 COMM_KB_INPUT；默认多目录合并。"""
    if cli is not None:
        p = cli.expanduser().resolve()
        return [p] if p.is_dir() else []
    multi = os.environ.get("COMM_KB_INPUTS", "").strip()
    if multi:
        out: list[Path] = []
        seen: set[str] = set()
        for part in multi.split(";"):
            part = part.strip()
            if not part:
                continue
            q = Path(part).expanduser().resolve()
            if q.is_dir():
                k = str(q)
                if k not in seen:
                    seen.add(k)
                    out.append(q)
        return out
    single = os.environ.get("COMM_KB_INPUT", "").strip()
    if single:
        p = Path(single).expanduser().resolve()
        return [p] if p.is_dir() else []
    return _default_comm_input_dirs()


ALLOWED_SOURCE_SUFFIX = frozenset({".docx", ".xlsx", ".xls"})


def discover_source_files(root: Path) -> list[Path]:
    """单层目录内所有支持的纪要文件（不递归子文件夹，避免误扫）。"""
    if not root.is_dir():
        return []
    found: list[Path] = []
    for p in sorted(root.iterdir(), key=lambda x: x.name.lower()):
        if p.is_file() and p.suffix.lower() in ALLOWED_SOURCE_SUFFIX:
            found.append(p)
    return found


def infer_archive_year(root: Path) -> str:
    m = re.search(r"(20\d{2})", root.name)
    return m.group(1) if m else "unknown"


def make_record_id(root: Path, path: Path) -> str:
    """跨年份、跨文件夹避免同名 id 冲突。"""
    stem = f"{root.name}-{path.stem}"
    return re.sub(r"\W+", "-", stem).strip("-").lower()[:100]


OUT_DIR = Path(__file__).resolve().parent / "out"

# (substring in filename, institution label) — longer / more specific first
INST_FILENAME_RULES: list[tuple[str, str]] = [
    ("久谦", "久谦咨询"),
    ("Morgan Stanley", "摩根士丹利"),
    ("摩根士丹利", "摩根士丹利"),
    ("JPMorgan", "摩根大通"),
    ("JP Morgan", "摩根大通"),
    ("JPM", "摩根大通"),
    ("摩根大通", "摩根大通"),
    ("星连资本", "星连资本"),
    ("浑瑾", "浑瑾投资"),
    ("老虎春秋", "老虎春秋基金"),
    ("纳德基金", "纳德基金"),
    ("前沿有理", "前沿有理"),
    ("工银理财", "工银理财"),
    ("中金财富", "中金财富"),
    ("中金线下", "中金公司"),
    ("中金", "中金公司"),
    ("中信建投", "中信建投"),
    ("银河证券", "银河证券"),
    ("招商银行", "招商银行"),
    ("招商交流", "招商证券"),
    ("光大银行", "光大银行"),
    ("中国银行", "中国银行"),
    ("德意志", "德意志银行"),
    ("花旗", "花旗银行"),
    ("方德证券", "方德证券"),
    ("李录", "其他"),
    ("巴菲特", "其他"),
    ("UBS", "瑞银"),
    ("各机构2026年经济展望", "多机构汇总"),
    ("美伊冲突", "其他"),
    ("中东石油危机", "其他"),
    ("日韩股市熔断", "其他"),
    ("港股资金流向分析-MS", "摩根士丹利"),
    ("美股市场情况", "其他"),
]

# 银行及银行理财类机构：资产类别统一突出「固收类」
BANK_FIXED_INCOME_INSTITUTIONS: frozenset[str] = frozenset(
    {
        "招商银行",
        "光大银行",
        "中国银行",
        "花旗银行",
        "德意志银行",
        "工银理财",
        "瑞银",
        "渣打银行",
        "汇丰银行",
        "星展银行",
    }
)

# 文件名中若出现以下片段则归为对应机构（按长度降序匹配，避免短词误吞）
_ORG_STEM_FRAGMENTS: tuple[str, ...] = tuple(
    sorted(
        {
            "法国巴黎银行",
            "摩根士丹利",
            "摩根大通",
            "中信建投",
            "中金公司",
            "中金财富",
            "招商证券",
            "招银国际",
            "招银理财",
            "银河证券",
            "德意志银行",
            "花旗银行",
            "中国银行",
            "招商银行",
            "光大银行",
            "工银理财",
            "方德证券",
            "久谦咨询",
            "浑瑾投资",
            "纳德基金",
            "前沿有理",
            "老虎春秋基金",
            "华泰证券",
            "中信证券",
            "易方达",
            "渣打银行",
            "汇丰银行",
            "北京银行",
            "工商银行",
            "华夏银行",
            "星展银行",
            "远东宏信",
            "宝盛银行",
            "嘉实基金",
            "九坤",
            "星连资本",
            "高盛",
            "琮碧秋实",
            "恩大基金",
            "方圆基金",
        },
        key=len,
        reverse=True,
    )
)

_RE_JPM_ORG = re.compile(r"j\s*p\s*m|jpmorgan|jp\s*morgan|摩根大通", re.I)
_RE_MS_ORG = re.compile(
    r"(?:"
    r"morgan\s*stanley|"
    r"摩根士丹利|"
    r"MS(?:全球|全|专栏|策略|宏观|美|港|A股|银行)|"
    r"(?:——|—|－|\s)MS\s*$|"
    r"(?:^|[\s\d年月])MS(?=全球|全|专栏|策略|宏观|202|\d)"
    r")",
    re.I,
)
_RE_GS_ORG = re.compile(r"高盛|goldman\s*sachs?|\bGS\.N\b", re.I)

# 文件名命中则归为「企业家访谈」（优先于机构规则）
_ENTREPRENEUR_FILENAME_HINTS: tuple[str, ...] = (
    "习近平与民营企业",
    "民营企业座谈会",
    "段永平",
    "蔡崇信",
    "王石",
    "黄铮",
    "黄峥",
    "胡猛",
)

_SPECIAL_INST_LABELS: frozenset[str] = frozenset({"其他", "多机构汇总", "企业家访谈"})

_LEGACY_INST_TO_OTHER: frozenset[str] = frozenset(
    {
        "其他机构",
        "未分类机构",
        "专题纪要",
        "市场分享纪要",
        "李录访谈",
        "巴菲特访谈",
    }
)

_KNOWN_CORE_INST: frozenset[str] = (
    frozenset(lab for _, lab in INST_FILENAME_RULES)
    | frozenset(_ORG_STEM_FRAGMENTS)
    | frozenset({"博时", "企业家访谈"})
)


def _strip_date_tokens_from_stem(stem: str) -> str:
    s = re.sub(r"20\d{6}", "", stem)
    s = re.sub(r"\d{8}", "", s)
    return re.sub(r"[\s_—-]+", " ", s).strip()


def _is_entrepreneur_related(stem: str, full: str) -> bool:
    blob = f"{stem}{full}"
    return any(h in blob for h in _ENTREPRENEUR_FILENAME_HINTS)


def _institution_from_stem_fragments(stem: str) -> str:
    if "方德" in stem:
        return "方德证券"
    if "渣打" in stem:
        return "渣打银行"
    if "汇丰" in stem:
        return "汇丰银行"
    if "星展" in stem:
        return "星展银行"
    if "博时" in stem:
        return "博时"
    for frag in _ORG_STEM_FRAGMENTS:
        if frag in stem:
            return frag
    return ""


def _bucket_title_as_other_institution(stem: str) -> str:
    if re.search(r"(美股|港股|A股|全球)市场波动(?:情况)?跟进", stem):
        return "其他"
    return ""


def _canonicalize_institution(label: str, stem: str, full: str) -> str:
    """合并同一机构的不同写法（如 JPM / 摩根大通），输出统一机构名。"""
    lab = re.sub(r"(20\d{6}|\d{8})$", "", label).strip(" -_—")
    hay = f"{stem}\n{full}\n{lab}"

    if _RE_JPM_ORG.search(hay) or "摩根大通" in hay:
        return "摩根大通"
    if _RE_MS_ORG.search(hay) or "摩根士丹利" in hay:
        return "摩根士丹利"
    if _RE_GS_ORG.search(hay):
        return "高盛"

    if lab == "瑞银 UBS" or "瑞银" in lab and "UBS" in lab:
        return "瑞银"

    if "博时" in hay or lab in ("博时基金", "博时国际"):
        return "博时"

    return lab if lab else label


def _finalize_institution_label(label: str) -> str:
    """博时合并；历史标签与无法识别的名称一律归为「其他」。"""
    if label in ("博时基金", "博时国际"):
        return "博时"
    if label in _LEGACY_INST_TO_OTHER:
        return "其他"
    if label in _SPECIAL_INST_LABELS:
        return label
    if label in _KNOWN_CORE_INST:
        return label
    return "其他"


def resolve_institution(filename: str) -> str:
    stem = Path(filename).stem
    full = Path(filename).name

    if _is_entrepreneur_related(stem, full):
        return _finalize_institution_label("企业家访谈")

    label = ""
    for sub, lab in INST_FILENAME_RULES:
        if sub in stem or sub in full:
            label = lab
            break

    if not label:
        label = _institution_from_stem_fragments(stem)
    if not label:
        label = _institution_from_stem_fragments(_strip_date_tokens_from_stem(stem))
    if not label and ("各机构" in stem or "各机构" in full):
        label = "多机构汇总"
    if not label:
        label = _bucket_title_as_other_institution(stem)
    if not label:
        cleaned = _strip_date_tokens_from_stem(stem)
        label = (cleaned[:40] if cleaned else "") or "其他"

    label = _canonicalize_institution(label, stem, full)
    return _finalize_institution_label(label)

ASSET_KEYWORDS: list[tuple[str, str]] = [
    ("加密货币|比特币|以太坊|BTC|ETH|数字资产", "加密"),
    ("港股|恒生|南向资金|港币", "港股"),
    ("A股|沪深|创业板|科创板|中证|上证|深证", "A股"),
    ("美股|纳斯达克|标普|道琼斯|NYSE|NASDAQ|Mag 7|七巨头", "美股"),
    ("美债|国债收益率|利率|降息|加息|联储|美联储|固收|信用债", "固收/利率"),
    ("原油|石油|OPEC|黄金|商品|铜价|大宗", "商品"),
    ("日元|韩元|外汇|汇率|美元", "外汇"),
    ("宏观|GDP|CPI|PCE|通胀|衰退|软着陆", "宏观"),
    ("日本股市|韩国股市|亚太", "亚太股市"),
]

# Known tickers / symbols (substring match; longer phrases first in scan order)
TICKER_HINTS = [
    "GS.N",
    "301265",
    "比亚迪",
    "Anthropic",
    "Applovin",
    "AppLovin",
    "微软",
    "MSFT",
    "NVDA",
    "AAPL",
    "GOOGL",
    "META",
    "TSLA",
    "AMZN",
]

# 常见个股/主体中文名（按长度降序匹配，减少短词误伤）
STOCK_NAME_HINTS: tuple[str, ...] = tuple(
    sorted(
        {
            "贵州茅台",
            "宁德时代",
            "中芯国际",
            "比亚迪",
            "隆基绿能",
            "中国平安",
            "招商银行",
            "腾讯控股",
            "阿里巴巴",
            "美团",
            "京东集团",
            "拼多多",
            "小米集团",
            "网易",
            "百度集团",
            "理想汽车",
            "蔚来",
            "小鹏汽车",
            "地平线机器人",
            "华新绿源",
            "英伟达",
            "台积电",
            "阿斯麦",
            "英特尔",
            "AMD",
            "博通",
            "甲骨文",
            "Meta Platforms",
            "Alphabet",
            "亚马逊",
            "苹果",
            "特斯拉",
            "微软",
            "谷歌",
            "阿里巴巴",
            "腾讯",
            "阿里",
            "百度",
            "京东",
            "小米",
            "网易",
            "美团",
            "拼多多",
            "理想",
            "蔚来",
            "小鹏",
            "地平线",
            "台积电",
            "英伟达",
            "高盛",
            "摩根士丹利",
            "摩根大通",
            "花旗银行",
            "德意志银行",
            "瑞银",
            "汇丰",
        },
        key=len,
        reverse=True,
    )
)


def _snippet_label(sentence: str) -> str:
    if re.search(r"风险|审慎|承压|减持|回避|波动|警惕|下行|担忧", sentence):
        return "风险/审慎"
    if re.search(r"看好|推荐|机会|增持|上行|超额|乐观|布局|首选|核心仓位", sentence):
        return "机会/观点"
    if re.search(r"同比|环比|亿元|万美元|%\s|增长|下滑|营收|利润|指引|估值|PE|PB", sentence):
        return "数据/场景"
    if re.search(r"若|如果|假设|情景|压力测试|敏感性", sentence):
        return "假设/情景"
    return "述评/背景"


# 与 _snippet_label 一致但略弱，避免与关键词层简单相加过度放大
_SNIPPET_LABEL_BIAS: dict[str, int] = {
    "风险/审慎": -1,
    "机会/观点": 1,
    "数据/场景": 0,
    "假设/情景": 0,
    "述评/背景": 0,
    "提及": 0,
}

# (正则, 权重)：仅匹配较长/较明确的表述；单字「机会」「布局」等易误判，已剔除
_BULLISH_PHRASES: list[tuple[str, int]] = [
    (r"看好|推荐(?:买入|增持)|建议增持|明确增持|增持(?:至|比例)|超配|加仓", 2),
    (r"估值修复|盈利改善|业绩超预期|戴维斯双击|上行空间|打开(?:上涨)?空间|显著受益", 2),
    (r"投资机会|配置价值|加大配置|乐观(?:看待|预期)|看多|有利(?:于|因素)", 1),
    (r"韧性较强|回暖|向好|修复(?:至|到)?|超额收益", 1),
    (r"\boverweight\b|\bbuy\b|\bbullish\b|\bupgrade\b", 2),
]

_BEARISH_PHRASES: list[tuple[str, int]] = [
    (r"减持|减配|清仓|看空|下调(?:目标|评级)|业绩承压|盈利下滑|大幅回撤", 2),
    (r"下行风险|泡沫(?:化|风险)?|诉讼|调查|疲弱|疲软|恶化|利空", 2),
    (r"谨慎(?:对待|配置)?|回避|警惕|波动加剧|亏损扩大", 1),
    (r"\bunderweight\b|\bsell\b|\bbearish\b|\bdowngrade\b|\bunderperform\b", 2),
]

# 缓和/否定式：对冲误把「风险」一律算空、或「不看好」算多
_BIAS_SOFTEN_BULL = re.compile(
    r"风险可控|风险有限|风险不大|风险(?:边际)?下降|风险缓释|风险消化|化解风险|监管风险已落地",
)
_BIAS_NEGATE_BULL = re.compile(
    r"不看好|未看好|难看好|缺乏(?:明确)?利好|没有(?:明显)?机会|机会有限|难现(?:单边)?上行",
)
_BIAS_SOFTEN_BEAR = re.compile(
    r"下滑(?:幅度)?(?:收窄|趋缓)|压力(?:缓解|减轻)|悲观预期(?:改善|修复)",
)


def _negated_before_match(text: str, start: int) -> bool:
    """匹配点前约 12 字内出现否定/缺乏，则该次命中视为无效（简单窗口）。"""
    if start <= 0:
        return False
    lo = max(0, start - 12)
    pre = text[lo:start]
    return bool(
        re.search(r"(不|未|无|没有|难以|并非|勿|别|缺乏)(?:\s|[\u4e00-\u9fff·]){0,5}$", pre)
    )


def _negated_before_english(text: str, start: int) -> bool:
    lo = max(0, start - 16)
    pre = text[lo:start].lower()
    return bool(re.search(r"\b(not|no|without|hardly|barely)\s+[\w\s]{0,12}$", pre))


def _accumulate_weighted_phrases(text: str, phrases: list[tuple[str, int]], sign: int) -> int:
    total = 0
    t = text or ""
    for pat, w in phrases:
        for m in re.finditer(pat, t, re.I):
            st = m.start()
            if m.group(0) and m.group(0)[0].isascii():
                if _negated_before_english(t, st):
                    continue
            else:
                if _negated_before_match(t, st):
                    continue
            total += sign * w
    return total


def _keyword_polarity_net(text: str) -> int:
    """短语级多空净分（已做否定过滤）；不含标签分。"""
    t = text or ""
    pos = _accumulate_weighted_phrases(t, _BULLISH_PHRASES, +1)
    neg = _accumulate_weighted_phrases(t, _BEARISH_PHRASES, -1)
    net = pos + neg
    if _BIAS_SOFTEN_BULL.search(t):
        net += 1
    if _BIAS_NEGATE_BULL.search(t):
        net -= 2
    if _BIAS_SOFTEN_BEAR.search(t):
        net += 1
    return net


def snippet_direction_score(text: str, label: str) -> int:
    """Positive = 偏多证据，负 = 偏空证据（规则，非投资建议）。"""
    score = _SNIPPET_LABEL_BIAS.get(label or "", 0)
    t = text or ""
    # 「提及」类占位句不参与关键词多空，避免无正文噪声
    if (label or "") == "提及" and "未截取" in t:
        return score
    score += _keyword_polarity_net(t)
    return score


def classify_bias_from_score(score: int) -> str:
    if score >= 3:
        return "偏多"
    if score <= -3:
        return "偏空"
    return "中性"


def enrich_snippets_bias(snips: list[dict[str, Any]]) -> None:
    for sn in snips:
        sc = snippet_direction_score(str(sn.get("text", "")), str(sn.get("label", "")))
        sn["bias_score"] = sc
        sn["bias"] = classify_bias_from_score(sc)


def summarize_mention_bias(snips: list[dict[str, Any]]) -> dict[str, Any]:
    """单文档内、单标的：综合多空倾向（要求证据强度，减少和稀泥误判）。"""
    enrich_snippets_bias(snips)
    counts = {"偏多": 0, "偏空": 0, "中性": 0}
    total = 0
    for sn in snips:
        b = str(sn.get("bias", "中性"))
        if b in counts:
            counts[b] += 1
        total += int(sn.get("bias_score", 0))

    # 仅「提及」占位：一律中性
    if snips and all(str(sn.get("label", "")) == "提及" for sn in snips):
        note = "仅提及标的，无可用观点句"
        return {
            "bias": "中性",
            "bias_score": 0,
            "bias_breakdown": counts,
            "bias_note": note,
        }

    strong_long = counts["偏多"] >= 2 or total >= 5
    strong_short = counts["偏空"] >= 2 or total <= -5
    if strong_long and not strong_short:
        overall = "偏多"
    elif strong_short and not strong_long:
        overall = "偏空"
    elif strong_long and strong_short:
        overall = "中性" if abs(total) < 3 else ("偏多" if total > 0 else "偏空")
    elif counts["偏多"] > counts["偏空"] and counts["偏多"] >= 1 and total >= 2:
        overall = "偏多"
    elif counts["偏空"] > counts["偏多"] and counts["偏空"] >= 1 and total <= -2:
        overall = "偏空"
    else:
        overall = "中性"
    note = f"摘录{counts['偏多']}多·{counts['偏空']}空·{counts['中性']}中·累计分{total}"
    return {
        "bias": overall,
        "bias_score": total,
        "bias_breakdown": counts,
        "bias_note": note,
    }


def _sentences_for_snippets(text: str) -> list[str]:
    parts = re.split(r"(?<=[。！？；\n])", text[:200000])
    out: list[str] = []
    for sent in parts:
        t = _WS.sub(" ", sent).strip()
        if len(t) >= 12:
            out.append(t)
    return out


def _phrase_in_sentence(sentence: str, phrase: str) -> bool:
    if not phrase:
        return False
    if phrase.isascii() and phrase.replace(".", "").isalnum():
        low = sentence.lower()
        p = phrase.lower()
        if "." in p:
            return p in low
        return re.search(r"(?<![A-Za-z0-9])" + re.escape(p) + r"(?![A-Za-z0-9])", low) is not None
    return phrase in sentence


def extract_snippets_for_phrase(
    text: str,
    phrase: str,
    max_snippets: int = 5,
    max_len: int = 220,
) -> list[dict[str, str]]:
    found: list[dict[str, str]] = []
    seen_sig: set[str] = set()
    for sent in _sentences_for_snippets(text):
        if not _phrase_in_sentence(sent, phrase):
            continue
        lab = _snippet_label(sent)
        st = sent if len(sent) <= max_len else sent[: max_len - 1] + "…"
        sig = re.sub(r"\s+", "", st)[:100]
        if sig in seen_sig:
            continue
        seen_sig.add(sig)
        found.append({"text": st, "label": lab})
        if len(found) >= max_snippets:
            break
    return found


def collect_stock_candidates(text: str, filename: str) -> list[tuple[str, str]]:
    """Unique (key, display) for stocks to index in one document."""
    blob = f"{filename}\n{text[:180000]}"
    pairs: list[tuple[str, str]] = []
    seen_key: set[str] = set()
    fund_six: set[str] = set()

    def add_key_display(key: str, display: str, force_display: bool = False) -> None:
        k = normalize_stock_key(key)
        if len(k) < 2:
            return
        disp = (display or k).strip() or k
        if k in seen_key:
            if force_display:
                for i, (ek, _) in enumerate(pairs):
                    if ek == k:
                        pairs[i] = (k, disp)
                        break
            return
        seen_key.add(k)
        pairs.append((k, disp))

    # 公司（代码）优先，便于展示「华新绿源（301265）」
    for m in re.finditer(
        r"([\u4e00-\u9fff·]{2,10})\s*[（(]\s*([03689]\d{5})\s*[）)]",
        blob,
    ):
        name, code = m.group(1), m.group(2)
        if code in STOCK_NUMERIC_BLOCKLIST:
            continue
        if is_cn_ashare_six(code):
            add_key_display(code, f"{name}（{code}）", force_display=True)
        else:
            fund_six.add(code)
    for m in re.finditer(r"(?<![0-9A-Za-z])([03689]\d{5})(?![0-9])", blob):
        code = m.group(1)
        if code in STOCK_NUMERIC_BLOCKLIST:
            continue
        if is_cn_ashare_six(code):
            add_key_display(code, code)
        else:
            fund_six.add(code)
    for m in re.finditer(r"\b([A-Z]{1,5})\.(N|O|K|B|A)\b", blob):
        add_key_display(f"{m.group(1)}.{m.group(2)}", f"{m.group(1)}.{m.group(2)}")
    for m in re.finditer(r"(?:HK|港股)\s*[:：]?\s*0?(\d{4,5})\b", blob, re.I):
        add_key_display(normalize_stock_key(f"HK{m.group(1)}"), normalize_stock_key(f"HK{m.group(1)}"))
    for name in STOCK_NAME_HINTS:
        if name in blob:
            add_key_display(name, name)
    for hint in TICKER_HINTS:
        if hint in blob:
            add_key_display(hint, hint)
    if fund_six:
        uniq = sorted(fund_six)
        add_key_display(
            FUND_INDEX_KEY,
            f"{FUND_INDEX_KEY}（{'、'.join(uniq)}）",
            force_display=True,
        )
    return pairs


def _phrases_for_stock(key: str, display: str) -> list[str]:
    """Try multiple surface forms to locate sentences (name vs code vs full display)."""
    seen: set[str] = set()
    out: list[str] = []
    if key == FUND_INDEX_KEY:
        for p in ("基金", "指数", "ETF", "指数基金", "基金和指数"):
            if len(p) >= 2 and p not in seen:
                seen.add(p)
                out.append(p)
        mx = re.search(r"（([^）]+)）", display)
        if mx:
            for part in re.split(r"[、,，\s]+", mx.group(1)):
                part = part.strip()
                if re.fullmatch(r"[03689]\d{5}", part) and len(part) >= 4 and part not in seen:
                    seen.add(part)
                    out.append(part)
        return out[:28]
    for p in (display, key):
        p = (p or "").strip()
        if len(p) >= 2 and p not in seen:
            seen.add(p)
            out.append(p)
    m = re.match(r"([\u4e00-\u9fff·]{2,10})[（(]\s*([03689]\d{5})\s*[）)]", display)
    if m:
        for p in (m.group(1), m.group(2)):
            if p not in seen:
                seen.add(p)
                out.append(p)
    return out


def build_stock_mentions(text: str, filename: str) -> list[dict[str, Any]]:
    """Per-file: stocks with labeled context snippets (scene / view)."""
    if not text.strip():
        return []
    cands = collect_stock_candidates(text, filename)
    mentions: list[dict[str, Any]] = []
    for key, display in cands:
        snips: list[dict[str, str]] = []
        for phrase in _phrases_for_stock(key, display):
            snips = extract_snippets_for_phrase(text, phrase, max_snippets=5)
            if snips:
                break
        if not snips:
            snips = [{"text": f"文中提及「{display}」，未截取到完整句子上下文。", "label": "提及"}]
        meta = summarize_mention_bias(snips)
        row: dict[str, Any] = {"key": key, "display": display, "snippets": snips}
        row.update(meta)
        mentions.append(row)
    return mentions


def build_stock_index(records: list[dict[str, Any]]) -> list[dict[str, Any]]:
    """Cross-document rollup: key -> aggregated snippets by file."""
    by_key: dict[str, dict[str, Any]] = {}
    for r in records:
        rid = r.get("id", "")
        for m in r.get("stock_mentions") or []:
            key = m.get("key", "")
            if not key:
                continue
            disp = m.get("display", key)
            if key not in by_key:
                by_key[key] = {"key": key, "display": disp, "docs": []}
            by_key[key]["docs"].append(
                {
                    "record_id": rid,
                    "file_name": r.get("file_name", ""),
                    "institution": r.get("institution", ""),
                    "archive_year": r.get("archive_year", ""),
                    "source_folder": r.get("source_folder", ""),
                    "date": r.get("date"),
                    "snippets": m.get("snippets") or [],
                    "bias": m.get("bias", "中性"),
                    "bias_score": m.get("bias_score", 0),
                    "bias_breakdown": m.get("bias_breakdown", {}),
                    "bias_note": m.get("bias_note", ""),
                }
            )
    out = list(by_key.values())
    for entry in out:
        bull = bear = neu = 0
        score_sum = 0
        for d in entry["docs"]:
            b = d.get("bias", "中性")
            score_sum += int(d.get("bias_score", 0))
            if b == "偏多":
                bull += 1
            elif b == "偏空":
                bear += 1
            else:
                neu += 1
        if bull > bear and bull >= 2:
            consensus = "偏多"
        elif bear > bull and bear >= 2:
            consensus = "偏空"
        elif score_sum >= 6:
            consensus = "偏多"
        elif score_sum <= -6:
            consensus = "偏空"
        elif bull > bear and bull == 1 and score_sum >= 3:
            consensus = "偏多"
        elif bear > bull and bear == 1 and score_sum <= -3:
            consensus = "偏空"
        else:
            consensus = "中性"
        entry["consensus_bias"] = consensus
        entry["consensus_bias_note"] = (
            f"跨{len(entry['docs'])}篇纪要：{bull}篇偏多·{bear}篇偏空·{neu}篇中性（累计分{score_sum}）"
        )
    out.sort(key=lambda x: (-len(x["docs"]), x.get("display") or x["key"]))
    return out


def read_docx(path: Path) -> str:
    from docx import Document

    doc = Document(str(path))
    parts: list[str] = []
    for p in doc.paragraphs:
        t = (p.text or "").strip()
        if t:
            parts.append(t)
    for table in doc.tables:
        for row in table.rows:
            cells = [(c.text or "").strip() for c in row.cells]
            if any(cells):
                parts.append(" | ".join(cells))
    return "\n".join(parts)


def read_xlsx(path: Path) -> str:
    from openpyxl import load_workbook

    wb = load_workbook(path, read_only=True, data_only=True)
    lines: list[str] = []
    for sheet in wb.worksheets:
        lines.append(f"## {sheet.title}")
        for row in sheet.iter_rows(values_only=True):
            cells = [str(c).strip() for c in row if c is not None and str(c).strip()]
            if cells:
                lines.append("\t".join(cells))
    wb.close()
    return "\n".join(lines)


def read_xls(path: Path) -> str:
    import xlrd

    book = xlrd.open_workbook(str(path))
    lines: list[str] = []
    for si in range(book.nsheets):
        sh = book.sheet_by_index(si)
        lines.append(f"## {sh.name}")
        for ri in range(sh.nrows):
            row = sh.row(ri)
            cells = []
            for c in row:
                v = c.value
                if v is None or v == "":
                    continue
                cells.append(str(v).strip())
            if cells:
                lines.append("\t".join(cells))
    return "\n".join(lines)


def read_any(path: Path) -> str:
    suf = path.suffix.lower()
    if suf == ".docx":
        return read_docx(path)
    if suf == ".xlsx":
        return read_xlsx(path)
    if suf == ".xls":
        return read_xls(path)
    return ""


def extract_date_from_name(name: str) -> str | None:
    m = re.search(r"(20\d{2})(\d{2})(\d{2})", name)
    if m:
        y, mo, d = m.groups()
        return f"{y}-{mo}-{d}"
    m2 = re.search(r"(20\d{2})年(\d{1,2})月(\d{1,2})日?", name)
    if m2:
        y, mo, d = m2.groups()
        return f"{y}-{int(mo):02d}-{int(d):02d}"
    return None


def extract_date_from_body(text: str) -> str | None:
    """Parse meeting date from opening paragraphs when filename has no YYYYMMDD."""
    if not text or not text.strip():
        return None
    head = text[:6000]
    m = re.search(r"(20\d{2})\s*年\s*(\d{1,2})\s*月\s*(\d{1,2})\s*日?", head)
    if m:
        y, mo, d = m.group(1), int(m.group(2)), int(m.group(3))
        return f"{y}-{mo:02d}-{d:02d}"
    m2 = re.search(r"(20\d{2})[-/.](\d{1,2})[-/.](\d{1,2})", head)
    if m2:
        y, mo, d = m2.group(1), int(m2.group(2)), int(m2.group(3))
        return f"{y}-{mo:02d}-{d:02d}"
    return None


def resolve_communicated_at(filename: str, text: str) -> str | None:
    return extract_date_from_name(filename) or extract_date_from_body(text)


KEY_PERSON_BLOCKLIST: frozenset[str] = frozenset(
    {
        "公司",
        "市场",
        "投资",
        "政策",
        "全球",
        "中国",
        "美国",
        "日本",
        "欧洲",
        "香港",
        "英国",
        "德国",
        "法国",
        "经济",
        "行业",
        "资产",
        "资金",
        "客户",
        "机构",
        "团队",
        "业务",
        "产品",
        "服务",
        "区域",
        "国内",
        "海外",
        "本次",
        "本场",
        "纪要",
        "策略",
        "展望",
        "报告",
        "分析",
        "研究",
        "数据",
        "建议",
        "风险",
        "收益",
        "核心",
        "整体",
        "结构",
        "层面",
        "方向",
        "机遇",
        "挑战",
        "主题",
        "内容",
        "部分",
        "章节",
        "问答",
        "交流",
        "会议",
        "论坛",
        "活动",
        "嘉宾",
        "主持",
        "各位",
        "大家",
        "我行",
        "贵司",
        "持有人",
        "投资人",
        "监管",
        "央行",
        "政府",
        "企业",
        "上市",
        "一级",
        "二级",
        "债券",
        "股票",
        "基金",
        "理财",
        "私人",
        "银行",
        "证券",
        "保险",
        "信托",
        "资管",
        "量化",
        "对冲",
        "多头",
        "空头",
        "宏观",
        "微观",
        "短期",
        "长期",
        "今年",
        "明年",
        "去年",
        "季度",
        "月度",
        "年度",
        "Inthat",
        "That",
        "With",
        "From",
    }
)

# 人物检索侧栏与筛选：仅保留以下人物，其余命中一律汇总为「其他」
KEY_PERSON_WHITELIST: frozenset[str] = frozenset(
    {"丁总", "罗总", "温总", "巴菲特", "李录", "段永平"}
)


def _clean_person_token(raw: str) -> str:
    s = re.sub(r"[\s\u3000]+", "", raw).strip(" ·、，,;；|｜")
    return s[:14] if s else ""


def extract_key_people(text: str, filename: str) -> list[str]:
    """Heuristic extraction of speaker / title-holder names (no LLM)."""
    blob = f"{filename}\n{text[:120000]}"
    seen: set[str] = set()
    candidates: list[str] = []

    def add(raw: str) -> None:
        name = _clean_person_token(raw)
        if len(name) < 2 or len(name) > 14:
            return
        if name in KEY_PERSON_BLOCKLIST:
            return
        if all("\u4e00" <= c <= "\u9fff" and c in "一二三四五六七八九十" for c in name):
            return
        if name not in seen:
            seen.add(name)
            candidates.append(name)

    stem = Path(filename).stem
    if "李录" in stem or "李录" in filename or "李录" in blob:
        add("李录")
    if "巴菲特" in stem or "巴菲特" in filename or "巴菲特" in blob:
        add("巴菲特")
    if "段永平" in stem or "段永平" in filename:
        add("段永平")
    if "段永平" in blob:
        add("段永平")
    for token in ("丁总", "罗总", "温总"):
        if token in filename:
            add(token)
    # 文件名中尊称多为单字+总（罗总、丁总）；避免「利罗总」吃掉「罗总」
    for m in re.finditer(r"([\u4e00-\u9fff])总(?=线下|线上|沟通|——|—|-)", filename):
        zong = m.group(1) + "总"
        if zong in KEY_PERSON_WHITELIST:
            add(zong)

    patterns = [
        r"(?:主讲|演讲人|报告人|分享人|讲师|嘉宾主持|主持人|沟通人)[:：\s]+([\u4e00-\u9fff·]{2,10})(?=[\s,，。\n]|$)",
        r"([\u4e00-\u9fff]{2,6})\s*(?:董事总经理|首席经济学家|首席策略师|首席国内策略分析师|首席分析师|策略分析师)",
        r"(?:分析师|研究员)[:：]\s*([\u4e00-\u9fff·]{2,10})",
        r"([\u4e00-\u9fff]{2,6})\s*董秘",
        r"(?:Mr\.|Ms\.|Dr\.)\s+([A-Z][a-z]+\s+[A-Z][a-z]+)\b",
    ]
    for pat in patterns:
        for m in re.finditer(pat, blob, re.I):
            add(m.group(1))

    head = "\n".join(blob.split("\n")[:35])
    for m in re.finditer(
        r"^([\u4e00-\u9fff·]{2,8})\s+([^\n]{0,24})$",
        head,
        re.M,
    ):
        tail = m.group(2)
        if re.search(r"董事|首席|总经|策略|分析|所长|合伙人", tail):
            add(m.group(1))

    out: list[str] = []
    seen2: set[str] = set()
    for name in candidates:
        if name in KEY_PERSON_WHITELIST:
            if name not in seen2:
                seen2.add(name)
                out.append(name)
    if any(n not in KEY_PERSON_WHITELIST for n in candidates):
        if "其他" not in out:
            out.append("其他")
    return out[:28]


def apply_bank_fixed_income_tags(institution: str, tags: list[str]) -> list[str]:
    """银行类沟通：将「固收/利率」并入并前置「固收类」。"""
    if institution not in BANK_FIXED_INCOME_INSTITUTIONS:
        return tags
    seen: set[str] = set()
    out: list[str] = []
    for t in tags:
        if t in ("固收/利率", "固收类"):
            continue
        if t not in seen:
            seen.add(t)
            out.append(t)
    return ["固收类"] + out


def infer_institution(filename: str) -> str:
    """文件名推断沟通机构；统一别名并尽量只保留机构名称。"""
    return resolve_institution(filename)


def infer_asset_tags(text: str) -> list[str]:
    tags: list[str] = []
    for pattern, label in ASSET_KEYWORDS:
        if re.search(pattern, text, re.I):
            tags.append(label)
    # de-dupe preserve order
    seen: set[str] = set()
    out: list[str] = []
    for t in tags:
        if t not in seen:
            seen.add(t)
            out.append(t)
    if not out:
        out.append("综合/未标注")
    return out


def extract_tickers(text: str, filename: str) -> list[str]:
    blob = f"{filename}\n{text[:120000]}"
    found: list[str] = []

    def add(s: str) -> None:
        s = s.strip()
        if len(s) < 2 or s in found:
            return
        found.append(s)

    # A-share 6-digit（排除黑名单；非沪深科创六位归入基金/指数，不在此列表重复出现）
    for m in re.finditer(r"(?<!\d)([03689]\d{5})(?!\d)", blob):
        code = m.group(1)
        if code in STOCK_NUMERIC_BLOCKLIST:
            continue
        if is_cn_ashare_six(code):
            add(code)

    # US symbols like GS.N, BRK.B
    for m in re.finditer(r"\b([A-Z]{1,5})\.(N|O|K|B|A)\b", blob):
        add(f"{m.group(1)}.{m.group(2)}")

    # HK 00700 / 0700（含 HK18000→HK1810 等别名）
    for m in re.finditer(r"(?:HK|港股|恒生)\s*[:：]?\s*0?(\d{4,5})\b", blob, re.I):
        add(normalize_stock_key(f"HK{m.group(1)}"))

    for hint in TICKER_HINTS:
        if hint in blob:
            add(hint)

    return found[:80]


def first_summary(text: str, limit: int = 600) -> str:
    t = re.sub(r"\s+", " ", text).strip()
    if len(t) <= limit:
        return t
    return t[: limit - 1] + "…"


_WS = re.compile(r"[\s\u3000]+")
# Line starts like "一、..." "1." "(1)" "• ..."
_BULLET_START = re.compile(
    r"^[\s\u3000]*(?:"
    r"[(（]?\s*(?:[一二三四五六七八九十百千万零]+|\d{1,3})\s*[)\）、\.:：．]\s*"
    r"|[-•·●▪▸►\*＊]\s*"
    r")"
)
# Inline Chinese enumeration: "一、xxx 二、yyy" (after line start / newline / clause end)
_INLINE_ENUM = re.compile(
    r"(?:^|[\n。；;])\s*"
    r"[(（]?\s*([一二三四五六七八九十百千万]+|\d{1,2})\s*[)\）、\.:：．]\s*"
    r"([^。；\n]{8,}?)(?=[。；\n]|$)"
)


def _trim_point(s: str, max_each: int) -> str:
    s = _WS.sub(" ", s).strip()
    s = re.sub(r"^[：:]\s*", "", s)
    if len(s) > max_each:
        s = s[: max_each - 1] + "…"
    return s


def _dedupe_points(points: list[str], min_len: int = 10) -> list[str]:
    seen: set[str] = set()
    out: list[str] = []
    for p in points:
        p = p.strip()
        if len(p) < min_len:
            continue
        sig = re.sub(r"\s+", "", p)[:120]
        if sig in seen:
            continue
        seen.add(sig)
        out.append(p)
    return out


def _score_sentence(s: str) -> int:
    if not s:
        return 0
    score = 0
    if re.search(r"\d", s):
        score += 2
    for kw in (
        "预计",
        "预测",
        "风险",
        "机会",
        "建议",
        "关注",
        "核心",
        "要点",
        "同比",
        "环比",
        "增速",
        "估值",
        "盈利",
        "政策",
        "利率",
        "降息",
        "通胀",
    ):
        if kw in s:
            score += 1
    if 20 <= len(s) <= 180:
        score += 1
    return score


def _sentences_from_blob(blob: str) -> list[str]:
    # Chinese-ish sentence split
    parts = re.split(r"(?<=[。！？；\n])", blob)
    out: list[str] = []
    for p in parts:
        t = _WS.sub(" ", p).strip()
        if len(t) >= 14:
            out.append(t)
    return out


def extract_key_points(
    text: str,
    max_points: int = 18,
    max_each: int = 240,
) -> list[str]:
    """Heuristic bullets from doc body (no LLM): lines, enumerations, then scored sentences."""
    if not text or not text.strip():
        return []
    raw = text.replace("\r\n", "\n").replace("\r", "\n").strip()
    collected: list[str] = []

    # Pass 1: each non-empty line (typical docx paragraph)
    for line in raw.split("\n"):
        line = line.strip()
        if not line or len(line) < 8:
            continue
        if line.startswith("##") and len(line) < 120:
            collected.append(_trim_point(line.lstrip("#").strip(), max_each))
            continue
        # Wide table row: split on |
        if line.count("|") >= 3 and len(line) > 350:
            cells = [c.strip() for c in re.split(r"\s*\|\s*", line) if len(c.strip()) > 10]
            collected.extend(cells[:10])
            continue
        # Long run-on line without list markers: split into sentences
        if len(line) > 420 and not _BULLET_START.match(line):
            subs = re.split(r"(?<=[。！？；])", line)
            for sub in subs:
                sub = sub.strip()
                if len(sub) >= 14:
                    collected.append(_trim_point(sub, max_each))
            continue
        if _BULLET_START.match(line):
            stripped = _BULLET_START.sub("", line, count=1).strip()
            if len(stripped) >= 8:
                collected.append(_trim_point(stripped, max_each))
            else:
                collected.append(_trim_point(line, max_each))
        else:
            collected.append(_trim_point(line, max_each))

    # Pass 2: inline "一、... 二、..." inside long chunks (single paragraph paste)
    extra: list[str] = []
    for chunk in list(collected):
        if len(chunk) < 120:
            continue
        for m in _INLINE_ENUM.finditer(chunk):
            body = (m.group(2) or "").strip()
            if len(body) >= 12:
                extra.append(_trim_point(body, max_each))
    collected.extend(extra)

    collected = _dedupe_points(collected, min_len=8)

    # Pass 3: pad with information-dense sentences (original order) until cap
    if len(collected) < max_points:
        sents = _sentences_from_blob(raw[:120000])
        for s in sents:
            if len(collected) >= max_points:
                break
            st = _trim_point(s, max_each)
            if len(st) < 16:
                continue
            if _score_sentence(st) < 2 and len(st) > 150:
                continue
            collected.append(st)
        collected = _dedupe_points(collected, min_len=8)

    # Cap count and length
    out: list[str] = []
    for p in collected:
        if len(out) >= max_points:
            break
        out.append(_trim_point(p, max_each))
    return out


def trim_title_echo(points: list[str], file_name: str) -> list[str]:
    """Drop leading item if it only repeats the file title (common in docx first line)."""
    if not points:
        return points
    stem = Path(file_name).stem.strip()
    if not stem:
        return points
    first = points[0].strip()
    if first == stem or first.replace("\u3000", " ") == stem.replace("\u3000", " "):
        return points[1:]
    if len(first) <= len(stem) + 2 and stem in first and len(first) < 40:
        return points[1:]
    return points


@dataclass
class Record:
    id: str
    file_name: str
    source_path: str
    archive_year: str
    source_folder: str
    institution: str
    date: str | None
    communicated_at: str | None
    asset_tags: list[str]
    tickers: list[str]
    stock_mentions: list[dict[str, Any]]
    key_people: list[str]
    key_points: list[str]
    summary: str
    char_count: int
    ingested_at: str


def main() -> None:
    import argparse

    ap = argparse.ArgumentParser()
    ap.add_argument("--input", type=Path, default=None, help="单个纪要目录（指定后不再默认合并多目录）")
    args = ap.parse_args()
    input_roots = resolve_input_dirs(args.input)
    if not input_roots:
        input_roots = []

    OUT_DIR.mkdir(parents=True, exist_ok=True)
    records: list[dict[str, Any]] = []
    errors: list[str] = []

    if not input_roots:
        errors.append(
            "未找到任何纪要目录。请将 2025交流记录 / 2026交流记录 放在桌面，"
            "或使用 --input / 环境变量 COMM_KB_INPUT 或 COMM_KB_INPUTS（分号分隔）。"
        )

    for root in input_roots:
        files = discover_source_files(root)
        if not files:
            errors.append(f"目录内无支持的文件（.docx/.xlsx/.xls）: {root}")
            continue
        year_lbl = infer_archive_year(root)
        folder_lbl = root.name
        for path in files:
            name = path.name
            rid = make_record_id(root, path)
            try:
                text = read_any(path)
            except Exception as e:  # noqa: BLE001
                errors.append(f"{path}: {e}")
                text = ""

            inst = infer_institution(name)
            communicated_at = resolve_communicated_at(name, text) if text else extract_date_from_name(name)
            if text:
                tags = apply_bank_fixed_income_tags(inst, infer_asset_tags(text))
            else:
                tags = ["解析失败/空文档"]
            tickers = extract_tickers(text, name) if text else []
            mentions = build_stock_mentions(text, name) if text else []
            people = extract_key_people(text, name) if text else []
            kpts = trim_title_echo(extract_key_points(text), name) if text else []
            rec = Record(
                id=rid,
                file_name=name,
                source_path=str(path.resolve()),
                archive_year=year_lbl,
                source_folder=folder_lbl,
                institution=inst,
                date=communicated_at,
                communicated_at=communicated_at,
                asset_tags=tags,
                tickers=tickers,
                stock_mentions=mentions,
                key_people=people,
                key_points=kpts,
                summary=first_summary(text) if text else "（无正文或解析失败）",
                char_count=len(text),
                ingested_at=datetime.now().isoformat(timespec="seconds"),
            )
            records.append(asdict(rec))

    # 稳定排序：年份降序，再日期降序，再文件名
    def _sort_key(rec: dict[str, Any]) -> tuple[Any, ...]:
        y = rec.get("archive_year") or ""
        yk = y if y and y != "unknown" else "0000"
        d = rec.get("date") or ""
        return (yk, d, rec.get("file_name", ""))

    records.sort(key=_sort_key, reverse=True)

    try:
        from quote_name_resolver import enrich_records_stock_displays

        enrich_records_stock_displays(records, OUT_DIR / "quote_name_cache.json")
    except Exception as e:  # noqa: BLE001
        errors.append(f"证券名称解析（雅虎/缓存）跳过: {e}")

    stock_index = build_stock_index(records)
    input_dirs_meta = [str(p.resolve()) for p in input_roots]
    by_year: dict[str, int] = {}
    for r in records:
        y = str(r.get("archive_year", "unknown"))
        by_year[y] = by_year.get(y, 0) + 1
    payload = {
        "meta": {
            "title": "2025–2026 交流纪要知识库",
            "input_dirs": input_dirs_meta,
            "input_dir": "; ".join(input_dirs_meta) if input_dirs_meta else "",
            "generated_at": datetime.now().isoformat(timespec="seconds"),
            "record_count": len(records),
            "stock_count": len(stock_index),
            "records_by_year": by_year,
            "errors": errors,
        },
        "stock_index": stock_index,
        "records": records,
    }

    json_path = OUT_DIR / "kb-data.json"
    json_path.write_text(json.dumps(payload, ensure_ascii=False, indent=2), encoding="utf-8")

    template_path = Path(__file__).resolve().parent / "panel.template.html"
    template = template_path.read_text(encoding="utf-8")
    embed = json.dumps(payload, ensure_ascii=False)
    embed = embed.replace("</script>", "<\\/script>")
    html = template.replace("__KB_JSON__", embed)
    (OUT_DIR / "panel.html").write_text(html, encoding="utf-8")

    print(f"Wrote {json_path}")
    print(f"Wrote {OUT_DIR / 'panel.html'}")
    if errors:
        print("Warnings:")
        for e in errors:
            print(" ", e)


if __name__ == "__main__":
    main()
